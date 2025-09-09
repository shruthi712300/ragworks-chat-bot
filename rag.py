
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
import PyPDF2, docx, pandas as pd

embeddings = HuggingFaceEmbeddings(model_name='all-MiniLM-L6-v2')
vector_store = None

def init_vector_store():
    global vector_store
    if not os.path.exists('vector_store'):
        os.makedirs('vector_store', exist_ok=True)
    if vector_store is None:
        vector_store = Chroma(persist_directory='vector_store', embedding_function=embeddings)

def read_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ''
            text += page_text + '\n'
    return text

def process_document(file_path, user_id, filename):
    if filename.lower().endswith('.pdf'):
        text = read_pdf(file_path)
    elif filename.lower().endswith('.docx'):
        import docx
        doc = docx.Document(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs])
    elif filename.lower().endswith('.csv'):
        df = pd.read_csv(file_path)
        text = df.to_string()
    elif filename.lower().endswith('.txt'):
        with open(file_path, 'r') as f:
            text = f.read()
    else:
        raise ValueError('Unsupported file format')
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    chunks = splitter.split_text(text)
    docs = []
    for i, chunk in enumerate(chunks):
        docs.append(Document(page_content=chunk, metadata={'source': filename, 'user_id': user_id, 'chunk': i}))
    global vector_store
    if vector_store is None:
        init_vector_store()
    vector_store.add_documents(docs)
    return len(docs)

def query_rag_system(query, user_id, k=5):
    global vector_store
    if vector_store is None:
        init_vector_store()
    # Note: different chroma builds may expect different filter formats
    try:
        results = vector_store.similarity_search(query, k=k, filter={'metadata.user_id': user_id})
    except Exception:
        # fallback without filter
        results = vector_store.similarity_search(query, k=k)
    context = '\n\n'.join([r.page_content for r in results])
    sources = list({r.metadata.get('source','unknown') for r in results})
    return context, sources
